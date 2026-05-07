"""Build REPORT.docx from REPORT.txt with figures and tables embedded.

Outputs the .docx next to REPORT.txt at the project root. Re-run after
editing REPORT.txt to regenerate.

Run from code/:
    python scripts/build_report_docx.py
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not installed. run:")
    print("    pip install python-docx")
    sys.exit(1)


ROOT = Path(__file__).resolve().parent.parent.parent
REPORT_TXT = ROOT / "REPORT.txt"
FIG_DIR = ROOT / "code" / "outputs" / "figures"
OUT_DOCX = ROOT / "REPORT.docx"

# theme: shades of red used in place of Word's default blue for headings
# and emphasised text.
RED_DARK = RGBColor(0x7E, 0x14, 0x16)   # near-crimson, used for Title
RED_MAIN = RGBColor(0xA0, 0x1B, 0x1D)   # heading 1
RED_MID  = RGBColor(0xC0, 0x39, 0x2B)   # heading 2
RED_SOFT = RGBColor(0xD9, 0x4A, 0x3C)   # heading 3 / table headers


FIGURES = {
    "4.2": (
        "training_curves.png",
        "Training and validation curves over 40 epochs. Best validation "
        "Dice 0.9436 reached mid-training.",
    ),
    "5.1": (
        "calibration_per_chip.png",
        "Per-image calibration ratio (mm^2 per pixel) across the 41 "
        "labeled images. Coefficient of variation 2.17%.",
    ),
    "5.2": (
        "sample_overlay_C12_Bg1_z1.png",
        "Predicted seedable mask (green) overlaid on a held-out chip "
        "(C12). Validation Dice 0.928 on this image.",
    ),
    "5.3": (
        "area_scatter.png",
        "Predicted vs ground-truth seedable area, averaged over the "
        "seven background variants per chip. MAE 3.37 mm^2, R^2 = 0.642.",
    ),
    "5.4": (
        "test1_per_chip.png",
        "Per-chip mean predicted area on Testing_dataset_1, with one-"
        "sigma error bars across background variants.",
    ),
    "5.5": (
        "area_scatter_comparison.png",
        "Scratch U-Net vs ImageNet-pretrained ResNet34 encoder, end-to-"
        "end predicted vs ground-truth area on the training set's 24 "
        "chips.",
    ),
}


# ----- block detectors ------------------------------------------------------

UNDERLINE_RX = re.compile(r"^[=\-]+\s*$")
SECTION_RX = re.compile(r"^(\d+)\.\s+([A-Z][A-Z\s]*)\s*$")
SUBSECTION_RX = re.compile(r"^(\d+\.\d+)\s+(.+?)\s*$")
BULLET_RX = re.compile(r"^\s+-\s+\S")
# indented "label: body" (lowercase label, e.g. "  scratch: ...")
LABELED_RX = re.compile(r"^\s+([a-z][a-z\-]*):\s+(.*)$")
# top-level "Label: body" (capitalised one-to-three-word label at column 0,
# used for paragraph callouts like "Headline: ..." or "Net effect: ...")
TOP_LABELED_RX = re.compile(r"^([A-Z][A-Za-z]+(?:\s[a-z]+){0,2}):\s+(\S.*)$")
STAT_RX = re.compile(r"^\s+(.+?)\s*=\s*(.+?)\s*$")


def is_table_separator(line):
    """Multi-column dash separator like '----   ----   ----'. Single
    full-line '---' / '===' header underlines are handled separately."""
    s = line.strip()
    if not s or not all(c in "- " for c in s):
        return False
    return bool(re.search(r"-{2,}\s+-{2,}", s))


# ----- emitters -------------------------------------------------------------

def emit_prose(lines, i, doc):
    buf = []
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            break
        if line.startswith(" ") or line.startswith("\t"):
            break
        if UNDERLINE_RX.match(line):
            break
        if SECTION_RX.match(line) or SUBSECTION_RX.match(line):
            break
        buf.append(line.strip())
        i += 1
    text = " ".join(s for s in buf if s)
    if not text:
        return i
    # check for a "Label: body" callout at column 0 so it gets the same
    # bold-red label treatment as indented labeled blocks (scratch / pretrained)
    m = TOP_LABELED_RX.match(text)
    if m:
        label = m.group(1)
        body = m.group(2)
        p = doc.add_paragraph()
        r1 = p.add_run(label + ": ")
        r1.bold = True
        r1.font.color.rgb = RED_MID
        p.add_run(body)
    else:
        doc.add_paragraph(text)
    return i


def emit_indented_prose(lines, i, doc):
    buf = []
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            break
        if not (line.startswith(" ") or line.startswith("\t")):
            break
        if UNDERLINE_RX.match(line):
            break
        if BULLET_RX.match(line) or LABELED_RX.match(line) or STAT_RX.match(line):
            break
        if i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            break
        buf.append(line.strip())
        i += 1
    text = " ".join(s for s in buf if s)
    if text:
        p = doc.add_paragraph(text)
        p.paragraph_format.left_indent = Inches(0.25)
    return i


def emit_bullets(lines, i, doc):
    while i < len(lines):
        line = lines[i]
        m = re.match(r"^\s+-\s+(.*)$", line)
        if not m:
            break
        first = m.group(1)
        cont = []
        j = i + 1
        while j < len(lines):
            nxt = lines[j]
            if not nxt.strip():
                break
            if re.match(r"^\s+-\s+", nxt):
                break
            if not (nxt.startswith(" ") or nxt.startswith("\t")):
                break
            cont.append(nxt.strip())
            j += 1
        text = " ".join([first] + cont)
        doc.add_paragraph(text, style="List Bullet")
        i = j
        # allow a single blank line between bullets in the same list
        if i < len(lines) and not lines[i].strip():
            k = i + 1
            if k < len(lines) and re.match(r"^\s+-\s+", lines[k]):
                i = k
                continue
        break
    return i


def emit_labeled(lines, i, doc):
    line = lines[i]
    m = LABELED_RX.match(line)
    if not m:
        return i + 1
    label = m.group(1)
    first = m.group(2).rstrip()
    cont = []
    j = i + 1
    while j < len(lines):
        nxt = lines[j]
        if not nxt.strip():
            break
        if LABELED_RX.match(nxt):
            break
        if not (nxt.startswith(" ") or nxt.startswith("\t")):
            break
        cont.append(nxt.strip())
        j += 1
    body = " ".join([first] + cont)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r1 = p.add_run(label + ": ")
    r1.bold = True
    r1.font.color.rgb = RED_MID
    p.add_run(body)
    return j


def emit_stats(lines, i, doc):
    rows = []
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            break
        m = STAT_RX.match(line)
        if not m:
            break
        rows.append((m.group(1).strip(), m.group(2).strip()))
        i += 1
    if rows:
        table = doc.add_table(rows=len(rows), cols=2)
        table.autofit = False
        for r, (k, v) in enumerate(rows):
            cell_k = table.cell(r, 0)
            cell_v = table.cell(r, 1)
            cell_k.text = k
            cell_v.text = v
            cell_k.width = Inches(2.5)
            cell_v.width = Inches(2.5)
            for para in cell_k.paragraphs + cell_v.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
        # add small space after the stats block
        doc.add_paragraph()
    return i


def emit_table(lines, i, doc):
    header_line = lines[i]
    sep_line = lines[i + 1]

    cols = []
    in_dash = False
    start = 0
    for k, ch in enumerate(sep_line):
        if ch == "-":
            if not in_dash:
                start = k
                in_dash = True
        else:
            if in_dash:
                cols.append((start, k))
                in_dash = False
    if in_dash:
        cols.append((start, len(sep_line)))

    if len(cols) < 2:
        return emit_indented_prose(lines, i, doc)

    def split_cells(line):
        cells = []
        for j, (s, _) in enumerate(cols):
            end = cols[j + 1][0] if j + 1 < len(cols) else max(len(line), cols[-1][1])
            cell = line[s:end] if s < len(line) else ""
            cells.append(cell.strip())
        return cells

    headers = split_cells(header_line)

    rows = []
    j = i + 2
    while j < len(lines):
        line = lines[j]
        if not line.strip():
            break
        if not (line.startswith(" ") or line.startswith("\t")):
            break
        if SECTION_RX.match(line) or SUBSECTION_RX.match(line):
            break
        if is_table_separator(line):
            break
        rows.append(split_cells(line))
        j += 1

    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    for c, h in enumerate(headers):
        cell = table.cell(0, c)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RED_SOFT
            run.font.size = Pt(10)
    for r, row in enumerate(rows):
        for c, txt in enumerate(row):
            if c < n_cols:
                cell = table.cell(r + 1, c)
                cell.text = txt
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return j


# ----- figure insertion -----------------------------------------------------

def insert_figure(key, doc):
    info = FIGURES.get(key)
    if info is None:
        return
    fname, caption = info
    p = FIG_DIR / fname
    if not p.exists():
        print(f"  warning: figure missing, skipping: {fname}")
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(p), width=Inches(5.5))
    cap = doc.add_paragraph(f"Figure: {caption}")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RED_MID


# ----- main -----------------------------------------------------------------

def apply_red_theme(doc):
    """Override Word's default blue heading color with shades of red."""
    style_colors = {
        "Title":     RED_DARK,
        "Heading 1": RED_MAIN,
        "Heading 2": RED_MID,
        "Heading 3": RED_SOFT,
    }
    styles = doc.styles
    for name, color in style_colors.items():
        try:
            style = styles[name]
        except KeyError:
            continue
        font = style.font
        font.color.rgb = color


def main():
    if not REPORT_TXT.exists():
        print(f"REPORT.txt not found at {REPORT_TXT}")
        sys.exit(1)

    doc = Document()
    apply_red_theme(doc)

    title = doc.add_heading(
        "Real-Time Microchip Segmentation and Usable Surface Area "
        "Quantification",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Project Report")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lines = REPORT_TXT.read_text(encoding="utf-8").splitlines()

    i = 0
    while i < len(lines) and not SECTION_RX.match(lines[i]):
        i += 1

    prev_key = None

    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue
        if UNDERLINE_RX.match(line):
            i += 1
            continue

        m_sec = SECTION_RX.match(line)
        if m_sec:
            if prev_key is not None:
                insert_figure(prev_key, doc)
            doc.add_heading(f"{m_sec.group(1)}. {m_sec.group(2).strip().title()}",
                            level=1)
            prev_key = m_sec.group(1)
            i += 1
            continue

        m_sub = SUBSECTION_RX.match(line)
        if m_sub:
            if prev_key is not None:
                insert_figure(prev_key, doc)
            doc.add_heading(f"{m_sub.group(1)} {m_sub.group(2)}", level=2)
            prev_key = m_sub.group(1)
            i += 1
            continue

        if line.startswith(" ") or line.startswith("\t"):
            if i + 1 < len(lines) and is_table_separator(lines[i + 1]):
                i = emit_table(lines, i, doc)
                continue
            if BULLET_RX.match(line):
                i = emit_bullets(lines, i, doc)
                continue
            if STAT_RX.match(line):
                i = emit_stats(lines, i, doc)
                continue
            if LABELED_RX.match(line):
                i = emit_labeled(lines, i, doc)
                continue
            i = emit_indented_prose(lines, i, doc)
            continue

        i = emit_prose(lines, i, doc)

    if prev_key is not None:
        insert_figure(prev_key, doc)

    doc.save(OUT_DOCX)
    print(f"wrote REPORT.docx  ({OUT_DOCX.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
